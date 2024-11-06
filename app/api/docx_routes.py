from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from docx import Document
from app.utils.pii_detection import detect_pii, detect_docType
from app.utils.docx_processing import redact_docx_content, process_docx_file
import io

router = APIRouter()

async def read_and_validate_docx(file: UploadFile):
    contents = await file.read()
    if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raise HTTPException(status_code=400, detail="Invalid file format. Only DOCX is supported.")
    return contents

def extract_text_from_docx(doc: Document) -> str:
    """Extracts full text from a DOCX document for processing."""
    return "\n".join([para.text for para in doc.paragraphs])

@router.post("/detect")
async def detect_pii_in_docx(file: UploadFile = File(...)):
    contents = await read_and_validate_docx(file)
    
    # Process document and detect PII
    doc = Document(io.BytesIO(contents))
    text = extract_text_from_docx(doc)
    pii_types, document_type = detect_pii(text)
    
    return {"document_type": document_type, "detected_pii": pii_types}

@router.post("/redact")
async def redact_pii_in_docx(file: UploadFile = File(...), pii_to_redact: str = Form("all")):
    contents = await read_and_validate_docx(file)
    
    # Load DOCX file into Document object
    doc = Document(io.BytesIO(contents))
    full_text = extract_text_from_docx(doc)
    document_type = detect_docType(full_text)
    
    # Redact based on selected PII types
    if pii_to_redact == "all":
        redacted_doc, redacted_texts = redact_docx_content(doc, document_type)
    else:
        pii_to_redact_list = [item.strip() for item in pii_to_redact.split(",")]
        redacted_doc, redacted_texts = process_docx_file(doc, document_type, pii_to_redact_list)
    
    # Save redacted document to a buffer for streaming response
    output = io.BytesIO()
    redacted_doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=redacted_doc.docx"}
    )
